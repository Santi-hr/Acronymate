import re
import docx
import lxml
from src.common import defines as dv
from src.common import configVars as cv
from src.common import pathHelpers
from src.cmdInterface import ansiColorHelper, cmdProgressBar, userCmdHandler


class DocxReader:
    """Class to process the input document to find its acronyms"""
    def __init__(self, acro_dict_handler):
        self.acro_dict_handler = acro_dict_handler
        self.document = None
        self.full_regex = ""

        self._doc_namespace = None

    def extract_acro_word(self, filepath, filename_overwrite=None):
        """Main function. Opens a docx file and extracts its acronyms

        :param filepath: Path string to a docx file
        :param acro_dict_handler: Acronym dictionary objects
        :param filename_overwrite: String that overwrites the loaded file name stored in the acro handler
        """
        # 1. Open file
        self.document = docx.Document(filepath)
        self._doc_namespace = self.document.element.nsmap
        self.acro_dict_handler.str_file = pathHelpers.get_filename_from_path(filepath)
        # Overwrite filename if needed. This is used with the word extension to keep only one temp file
        if filename_overwrite:
            self.acro_dict_handler.str_file = filename_overwrite

        userCmdHandler.print_acronym_search_start()

        # 2. Set regex expression
        # Includes acronyms or abbreviates that do not match with the main regex from the DB or the acronym table
        # on the document (Ej: ExCOMMS, JdP)
        main_regex = cv.config_regex_acro_find #Todo: añadir test de documento sin tabla de acros

        brute_regex_list = []

        if cv.config_use_acro_from_doc_table:  # Special acronyms from the current document acronym
            self.__search_and_process_acronym_table()
            brute_regex_list = [acro_key for acro_key in self.acro_dict_handler.acros_doc_table.keys()
                                if not re.fullmatch(cv.config_regex_acro_find, acro_key)]

        if cv.config_use_non_matching_acro_from_db:  # Special acronyms from DB
            brute_regex_list += [acro_key for acro_key in self.acro_dict_handler.obj_db.list_no_regex
                                 if acro_key not in brute_regex_list]

        # All these words are combined into one big regex. It speeds up drastically the search as only one pass is needed
        # fixme: probably not the best use of regex, but if not many "special" acronyms it will work flawlessly.
        brute_regex_list_escaped = [re.escape(acro_key) for acro_key in brute_regex_list]
        brute_regex = r'|'.join(brute_regex_list_escaped)

        if brute_regex != "":
            # Add first brute search. This prevents from finding twice acronyms like ExCOMMS (ExCOMMS and COMMS)
            main_regex = r'\b(' + brute_regex + r')(?![\wÁÉÍÓÚ])|' + main_regex
            # Last part fixes abbreviates. \b does not match '. ' as both are non alphanumeric
            # Alternate: r'\b('+second_regex[:-1]+r')(\b|(?=\W))'
        self.full_regex = main_regex

        # 3. Search acronyms in the document using the set regex
        self.__extract_acro_from_all_paragraphs()
        self.__extract_acro_from_all_tables()
        self.__extract_acro_from_all_sections()

    def __extract_acro_from_all_paragraphs(self):
        """Iterates through all paragraph blocks of the document and stores all acronyms found"""
        # Iterate trough all paragraphs
        obj_progress_bar = cmdProgressBar.CmdProgressBar(len(self.document.paragraphs),
                                                         userCmdHandler.get_translated_str_paragraphs())
        for i, paragraph in enumerate(self.document.paragraphs):
            self.__extract_acro_from_paragraph(paragraph)
            obj_progress_bar.update(i + 1)

    def __extract_acro_from_all_tables(self):
        """Iterates through all document tables and stores all acronyms found"""
        # Iterate trough all table objects
        obj_progress_bar = cmdProgressBar.CmdProgressBar(len(self.document.tables),
                                                         userCmdHandler.get_translated_str_tables())
        for i, table in enumerate(self.document.tables):
            self.__extract_acro_from_table(table)
            obj_progress_bar.update(i + 1)

    def __extract_acro_from_all_sections(self):
        """Iterates through all sections blocks of the document and stores all acronyms found"""
        # Iterate trough all sections header and footers
        obj_progress_bar = cmdProgressBar.CmdProgressBar(len(self.document.sections),
                                                         userCmdHandler.get_translated_str_sections())
        for i, section in enumerate(self.document.sections):
            # Headers
            for paragraph in section.header.paragraphs:
                self.__extract_acro_from_paragraph(paragraph)
            for table in section.header.tables:
                self.__extract_acro_from_table(table)

            # Footers
            for paragraph in section.footer.paragraphs:
                self.__extract_acro_from_paragraph(paragraph)
            for table in section.footer.tables:
                self.__extract_acro_from_table(table)

            obj_progress_bar.update(i + 1)

    def __extract_acro_from_paragraph(self, paragraph): #todo: doc
        str_accepted_text = self.accepted_text(paragraph, paragraph._p.xml, self._doc_namespace)
        self.__extract_acro_from_str(str_accepted_text)

    def __extract_acro_from_table(self, table): #todo: doc
        for j, row in enumerate(table.rows):
            row_cell_list = [self.accepted_text(cell, cell._tc.xml, self._doc_namespace) for cell in row.cells]
            row_text = dv.define_tb_col_separator.join(row_cell_list)

            # Do not process acronym table as found acronyms
            if j == 0:
                if self.__is_acronym_table_header(row_text):
                    break

            self.__extract_acro_from_str(row_text)

    def __extract_acro_from_str(self, str_in_raw):
        """Finds acronyms in a text string and stores them into a dictionary

        :param str_in_raw: Input text string
        """
        # 1. Find Acronyms as regex matches of groups of N Capital Leters. The regex can be changed for other uses
        regex_in = self.full_regex

        # Line breaks are removed to reduce space used when outputting the context string to console
        str_in = str_in_raw.replace('\n', dv.define_new_line_separator)
        re_results = re.finditer(regex_in, str_in)

        if re_results:
            # 2. Iterate trough all matches (There could be multiple acronyms in a paragraph)
            for re_result in re_results:
                # 3. Save context for user analysis (Only part of the string)
                context_width = 200
                # Ensure the context width is achieved
                prev_context = context_width / 2
                next_context = context_width / 2
                # Check if there are unused characters at front/back after adding/subtracting the context. If so, add
                # them to the opposite side.
                char_diff_ini = re_result.start(0) - context_width / 2
                if char_diff_ini < 0:
                    next_context -= char_diff_ini  # Addition, double negative
                char_diff_end = len(str_in) - (re_result.start(0) + context_width / 2)
                if char_diff_end < 0:
                    prev_context -= char_diff_end  # Addition, double negative

                context_substr_ini = int(max(re_result.start(0) - prev_context, 0))
                context_substr_end = int(min(re_result.start(0) + next_context, len(str_in)))

                # Context with acronym remarked using ANSI Color Codes
                str_context = ansiColorHelper.highlight_substr(
                    str_in[context_substr_ini:context_substr_end],
                    (re_result.span()[0] - context_substr_ini, re_result.span()[1] - context_substr_ini),
                    ansiColorHelper.AnsiColorCode.CYAN)

                # 4. Create empty dict if acronym is first found
                self.acro_dict_handler.add_acronym_found(re_result.group(0), str_context)

    def __search_and_process_acronym_table(self):
        """ Iterates through all document tables and stores all acronyms found. Processes the acro-table if found"""
        for table in self.document.tables:
            row_cell_list = [self.accepted_text(cell, cell._tc.xml, self._doc_namespace) for cell in table.rows[0].cells]
            row_text = dv.define_tb_col_separator.join(row_cell_list)

            if self.__is_acronym_table_header(row_text):
                self.__process_acronym_table(table)
                break  # Do not process other acronym tables. Only one expected

    def __is_acronym_table_header(self, row_input):
        """Checks if the string matches any of the document acronyms table header defined

        :param row_input: Header table string
        :return: True if any match is found
        """
        flag_return = False
        # Iterate trough all acronym headers
        for header_columns in cv.config_acronym_table_headers:
            header = dv.define_tb_col_separator.join(header_columns)
            if row_input == header:
                flag_return = True
                break

        return flag_return

    def __process_acronym_table(self, acro_table):
        """Processes the document acronym table to extract already defined acronyms. Currently the format is set as:
        Acronym | Definitions as "original text (Translated)" separated by line breaks

        :param acro_table: Python-docx table object
        """
        # This function only works if the following format is used:
        # Acronym | Definitions as "original text (Translated)" separated by line breaks
        # When a table is found with a different number of rows it is skiped
        # Todo: Allow for custom tables via configure file
        # Fixme: Prevent crashing if table has merged lines
        if not self.acro_dict_handler.flag_doc_table_processed and len(acro_table.rows[0].cells) == 2:
            for row in acro_table.rows[1:]:  # Skipping header (Row 1)
                # Get raw data
                acronym = row.cells[0].text.strip()
                definitions = row.cells[1].text.split('\n')

                if acronym == "":  # Skip empty rows
                    continue

                # Process definitions
                for raw_def in definitions:
                    main_def = raw_def.strip()
                    trans_def = ""

                    if ')' in main_def:
                        # Finds the opening and closing parenthesis of the translated part to slice the definition
                        opening_par = -1
                        closing_par = main_def.rfind(')')
                        groups_open = 1  # Group = (). Assumes proper closing of parenthesis
                        for idx in range(closing_par - 1, 0, -1):
                            if main_def[idx] == ')':
                                groups_open += 1
                            elif main_def[idx] == '(':
                                groups_open -= 1
                            if groups_open == 0:
                                opening_par = idx
                                break
                        trans_def = main_def[opening_par + 1:closing_par].strip()
                        main_def = main_def[:opening_par - 1].strip()

                    self.acro_dict_handler.add_acronym_doc_table(acronym, main_def, trans_def)
        self.acro_dict_handler.flag_doc_table_processed = True

    @staticmethod
    def accepted_text(docx_elem, docx_elem_xml, doc_nsmap): #todo: documentar. Pasar solo docx_elem?
        str_accepted_text = ""
        if "w:ins" in docx_elem_xml:
            target = lxml.etree.XML(docx_elem_xml)
            # Search all paragraphs in the XML
            for paragraph in target.xpath('//w:p', namespaces=doc_nsmap):
                # Search all runs and inserted runs (Track Changes). Deleted runs (w:del) are skipped
                for text_run in paragraph.xpath('w:r | w:ins/w:r', namespaces=doc_nsmap):
                    # Handle linebreaks first. It seems that appear in their own run or before text
                    if text_run.xpath('w:br', namespaces=doc_nsmap):
                        str_accepted_text += '\n'
                    for text_tag in text_run.xpath('w:t', namespaces=doc_nsmap):
                        str_accepted_text += text_tag.text
                str_accepted_text += '\n'  # Append new line between paragraphs (There can be multiple pgphs in tables)
            str_accepted_text = str_accepted_text[:-1]  # Remove last paragraph new line char
        else:
            # Use python docx when possible as it has the text already extracted
            str_accepted_text = docx_elem.text
        return str_accepted_text
