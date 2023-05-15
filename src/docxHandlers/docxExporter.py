import docx
import time
from src.common import defines as dv
from src.common import configVars as cv
from src.common import stringHelpers as strHlprs
from src.cmdInterface import cmdProgressBar


class DocxExporter:
    """Class to generate the output acronym document"""
    def __init__(self, acro_dict_handler):
        self.acro_dict_handler = acro_dict_handler
        self.document = None
        self.obj_normal_style = None
        self.obj_header_style = None
        self.time_paragraph = 0

    def generate_output_docx(self, time_begin_acronymate):
        """Generates a docx document with the user accepted acronyms

        :param time_begin_acronymate: time.monotonic() of the program start. Used in the stats print
        :return: python-docx document object
        """
        # --- Initialize ---
        print(_("Generando Word con tabla de acrónimos"))
        obj_progress_bar = cmdProgressBar.CmdProgressBar(len(self.acro_dict_handler.acros_output.keys()), _("Acrónimos"))

        self.document = docx.Document()

        # --- Define styles ---
        self._define_styles()

        # --- Footer ---
        self._add_footer()

        # --- Text before table (Statistics) ---
        self._add_summary_paragraphs()

        # --- Acronym Table ---
        self._add_acronym_table(obj_progress_bar)

        # --- Fill elapsed time to statistics ---
        self._complete_time_paragraph(time_begin_acronymate)

        return self.document

    def _define_styles(self):
        """Defines the styles used in the document"""
        self.obj_normal_style = self.document.styles['Normal']
        font = self.obj_normal_style.font
        font.name = cv.config_output_font
        font.size = docx.shared.Pt(cv.config_output_font_size)
        paragraph_format = self.obj_normal_style.paragraph_format
        paragraph_format.space_after = docx.shared.Pt(2)
        paragraph_format.space_before = docx.shared.Pt(2)

        self.obj_header_style = self.document.styles.add_style('HeaderStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        hdr_font = self.obj_header_style.font
        hdr_font.size = docx.shared.Pt(cv.config_output_font_size_header)
        hdr_font.name = cv.config_output_font
        paragraph_format = self.obj_header_style.paragraph_format
        paragraph_format.space_after = docx.shared.Pt(2)
        paragraph_format.space_before = docx.shared.Pt(2)

    def _add_footer(self):
        """Adds the footer of the document"""
        # Only one section in document
        section_footer = self.document.sections[0].footer
        section_footer.paragraphs[0].text = "ACRONYMATE " + dv.define_acronymate_version + " - SAHR Projects 2023"
        section_footer.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

    def _add_summary_paragraphs(self):
        """Adds a brief summary before the acronyms table"""
        self.document.add_paragraph(_('Archivo procesado: %s') % self.acro_dict_handler.str_file)
        self.document.add_paragraph(_('Acrónimos extraídos: %d') % len(self.acro_dict_handler.acros_output.keys()))
        self.document.add_paragraph('Tiempo')
        self.time_paragraph = len(self.document.paragraphs) - 1  # Save paragraph to fill later with a precise time
        self.document.add_paragraph('')  # Spacer

    def _complete_time_paragraph(self, time_begin_acronymate):
        """Completes the elapsed time paragraph with hours minutes and seconds since the beginning of the program"""
        elapsed_f_sec = time.monotonic() - time_begin_acronymate
        hours, minutes = 0, 0
        while elapsed_f_sec >= 3600:
            elapsed_f_sec -= 3600
            hours += 1
        while elapsed_f_sec >= 60:
            elapsed_f_sec -= 60
            minutes += 1
        seconds = int(elapsed_f_sec)
        msec = int((elapsed_f_sec - seconds)*100)

        self.document.paragraphs[self.time_paragraph].text = \
            _('Tiempo transcurrido: %s:%s:%s.%d') % \
            (str(hours).zfill(2), str(minutes).zfill(2), str(seconds).zfill(2), msec)

    def _add_acronym_table(self, obj_progress_bar):
        """Adds the acronym table"""
        table = self.document.add_table(rows=1, cols=2)
        # In case we would like to use some styling: table.obj_normal_style = 'LightShading-Accent1'
        table.autofit = True
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = _('Acrónimo')
        hdr_cells[1].text = _('Definición')
        for i in range(len(hdr_cells)):
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].style = self.obj_header_style
        # Set parameter in cells for auto width
        for cell in table.rows[0].cells:
            cell._tc.tcPr.tcW.type = 'auto'

        for n_acro, acro in enumerate(sorted(self.acro_dict_handler.acros_output.keys(), key=strHlprs.acro_ordering)):
            row_cells = table.add_row().cells
            row_cells[0].text = acro
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            for i in range(len(self.acro_dict_handler.acros_output[acro]['Def'])):
                if i > 0:
                    row_cells[1].add_paragraph()
                run_main_def = row_cells[1].paragraphs[i].add_run(
                    self.acro_dict_handler.acros_output[acro]['Def'][i]['Main'])
                if 'Translation' in self.acro_dict_handler.acros_output[acro]['Def'][i]:
                    run_main_def.italic = True
                    row_cells[1].paragraphs[i].add_run(
                        " (" + self.acro_dict_handler.acros_output[acro]['Def'][i]['Translation'] + ")")
            # Set parameter in cells for auto width
            for cell in row_cells:
                cell._tc.tcPr.tcW.type = 'auto'

            obj_progress_bar.update(n_acro+1)


def save_document(filepath_output, document):
    """Static function that saves an python-docx object

    :param filepath_output: Desired path
    :param document: python-docx document object
    """
    document.save(filepath_output)
